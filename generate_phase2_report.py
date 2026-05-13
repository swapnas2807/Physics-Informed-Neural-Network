import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Phase 2: Deep Learning Models for Forward & Inverse Prediction', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1) Previous MC Dropout
doc.add_heading('1. Limitations of Previous MC Dropout', level=1)
p = doc.add_paragraph()
p.add_run('Previously, Monte Carlo (MC) Dropout was used during gradient descent for inverse prediction, but it did not work well. The primary issue was that it projected ').bold = False
p.add_run('Epistemic Model Uncertainty').bold = True
p.add_run(' (the AI\'s lack of knowledge) directly onto the physical geometry parameters (Widths). Since physical dimensions are rigid layouts rather than probabilistic features, this resulted in unrealistic and physically invalid uncertainty bounds. The model used its own lack of confidence to incorrectly define structural bounds, rather than properly quantifying its forward prediction confidence.')

# 2) One to many mapping problem
doc.add_heading('2. The One-To-Many Mapping Problem', level=1)
p = doc.add_paragraph()
p.add_run('In a forward model (predicting Performance from Widths), the mapping is a strict One-to-One mathematical function. However, the inverse process (recovering Widths from a target Performance) suffers from a "One-To-Many" mapping problem. Multiple, entirely distinct geometric circuit configurations can yield the exact same electrical performance metrics. Traditional Generative AI models (like cVAE) failed here due to Mode Collapse—forcing a non-physical 1-to-1 shortcut map rather than exploring the diverse valid topology space.')

# 3) UQ in Forward vs Inverse
doc.add_heading('3. Uncertainty Quantification (UQ): Forward vs. Inverse', level=1)
p = doc.add_paragraph()
p.add_run('True Uncertainty Quantification fundamentally differs depending on directionality:\n').bold = True
p.add_run('• Forward Pass UQ (Epistemic Uncertainty): ').bold = True
p.add_run('This captures the model’s confidence over its performance given fixed widths. It acts as a strict guardrail to prevent hallucination in Out-of-Distribution (OOD) spaces, e.g., Outputting Gain as "50dB ± 1.5dB".\n')
p.add_run('• Inverse Output "UQ" (Design Tolerance): ').bold = True
p.add_run('This is not actually model uncertainty, but rather a representation of geographic "design flexibility" or margin. It defines how precisely a parameter must be hit to maintain valid performance across different valid topologies, e.g., Outputting Width as "10µm ± 0.5µm".\n')

# Diagram
diagram_path = r'c:\Users\LENOVO\Desktop\Physics Informed Neural Networks For EDA\uq_comparison_diagram.png'
if os.path.exists(diagram_path):
    doc.add_picture(diagram_path, width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 4) Objectives
doc.add_heading('4. Objectives for Phase 2', level=1)
p = doc.add_paragraph('Based on the problems discovered, the objectives were:')
p.style = 'List Bullet'
p2 = doc.add_paragraph('Solve the One-To-Many mapping problem using two distinct algorithmic approaches: Generative Conditional Variational Autoencoder (cVAE) and Multi-Start Parallel Optimization.')
p2.style = 'List Bullet'
p3 = doc.add_paragraph('Establish a scientifically sound Uncertainty Quantification system for both Forward (Model Confidence) and Inverse (Geometric Margins).')
p3.style = 'List Bullet'
p4 = doc.add_paragraph('Build an end-to-end interactive UI for seamless VLSI design space exploration.')
p4.style = 'List Bullet'

# 5) Work Done
doc.add_heading('5. Work Done', level=1)
p = doc.add_paragraph()
p.add_run('To achieve these objectives, the following implementation tasks were completed:\n').bold = True
p.add_run('• Benchmarked Architecture: ').bold = True
p.add_run('Both the cVAE and the Multi-Start approaches were fully built and benchmarked via an Inverse Evaluation Suite.\n')
p.add_run('• Adopted 1000-Start Method: ').bold = True
p.add_run('Based on severe Mode Collapse in the cVAE, it was abandoned. We selected the Multi-Start Global Optimizer, deploying 1,000 parallel Adam gradient descents initialized with stochastic noise to natively discover the entire solution space manifold.\n')
p.add_run('• Deep Ensemble for Forward UQ: ').bold = True
p.add_run('To prevent the standard PINN from actively hallucinating wrong results, we created a Deep Ensemble of 5 independently trained PINNs (different weight initializations). Their variance serves as Epistemic Uncertainty inside the optimizer loss function.\n')
p.add_run('• K-Means Standard Deviation for Width Bounds: ').bold = True
p.add_run('After running 1,000 starts, solutions are grouped via K-Means clustering. The standard deviation across each isolated cluster was adopted as the geometric margin representing layout flexibility (±).\n')
p.add_run('• User Interface (UI): ').bold = True
p.add_run('Built a Streamlit-based UI integrating Target input tools, the Multi-Start logic under the hood, and interactive scatter maps displaying clustered geometric topologies.')

# 6) Results
doc.add_heading('6. Results and Validation', level=1)

p = doc.add_paragraph()
p.add_run('Selection of the 1000-Start Optimizer: ').bold = True
p.add_run('The cVAE failed with a Diversity Score of 0.15µm, while the 1000-Start Optimizer achieved a Diversity Score of 13.56µm. It resolved distinct geometric families perfectly and successfully avoided the hallucination loophole by leveraging gradient traversal through multiple solution modes.')

p2 = doc.add_paragraph()
p2.add_run('Selection of the Deep Ensemble: ').bold = True
p2.add_run('Unlike MC Dropout, training an ensemble of 5 isolated PINNs stabilized backpropagation. Applying an Epistemic UQ Guardrail effectively prevented optimization in uncertain spaces, punishing the loss map dynamically based on prediction variance.')

# NEW: Plot 5 for UQ Calibration Guardrail
plot5 = r'c:\Users\LENOVO\Desktop\Physics Informed Neural Networks For EDA\Inverse Validation\plot5_uq_calibration.png'
if os.path.exists(plot5):
    p_img_uq = doc.add_paragraph('Fig: UQ Calibration showing the Epistemic Guardrail actively trapping uncertain/OOD spaces by correlating prediction error with Model Uncertainty.')
    p_img_uq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(plot5, width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph()
p3.add_run('K-Means for Standard Deviation: ').bold = True
p3.add_run('Using K-Means provided contextually aware "design margins" rather than monolithic structural bounds. Small standard deviations map seamlessly to critical sizing paths, whilst large deviations identify components with vast layout flexibility.')

# NEW: Plot 8 for Width Margin/Tolerance
plot8 = r'c:\Users\LENOVO\Desktop\Physics Informed Neural Networks For EDA\Inverse Validation\plot8_width_comparison.png'
if os.path.exists(plot8):
    p_img_w = doc.add_paragraph('Fig: Width Tolerance Spreads across K-Means clusters. Shows exactly which constraints are mathematically locked vs flexible.')
    p_img_w.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(plot8, width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

p4 = doc.add_paragraph()
p4.add_run('Validation of Inverse Technique: ').bold = True
p4.add_run('Running an automated validation suite across a held-out test set proved 100% Top-1 Yield with an average prediction error of just 2.14% (well within engineering tolerance limits).')

plot6 = r'c:\Users\LENOVO\Desktop\Physics Informed Neural Networks For EDA\Inverse Validation\plot6_design_diversity.png'
if os.path.exists(plot6):
    p_img = doc.add_paragraph('Fig: K-Means Spatial Layout Diversity Visualization (Resolving One-To-Many Mapping)')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(plot6, width=Inches(5.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

plot10 = r'c:\Users\LENOVO\Desktop\Physics Informed Neural Networks For EDA\Inverse Validation\plot10_summary_dashboard.png'
if os.path.exists(plot10):
    p_img2 = doc.add_paragraph('Fig: Complete Validation Summary Dashboard (Accuracy & Error Distribution)')
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(plot10, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'c:\Users\LENOVO\Desktop\Physics Informed Neural Networks For EDA\Phase_2_Documentation.docx')
print("Document saved successfully!")
